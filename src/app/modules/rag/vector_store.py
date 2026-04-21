from pathlib import Path
from app.modules.rag.process.splitter import SplitterFactory
from app.utils.logger_handler import get_logger
from langchain_chroma import Chroma
from app.shared.exception.file import FileTypeError
from app.core.config import settings
from app.core.llm.factory import get_embedding_model
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import DashScopeEmbeddings
from app.utils.file import check_md5_hex, save_md5_hex
from app.utils.file_handler import pdf_loader, text_loader, listdir_with_allowed_types,get_file_md5_hex
from langchain_core.documents import Document
from docx import Document as DocxDocument
logger = get_logger(__name__)
class VectorStoreService:
    def __init__(self):
        self.vector_store = Chroma(
            persist_directory=settings.chroma.persist_directory,
            collection_name=settings.chroma.collection_name,
            embedding_function=DashScopeEmbeddings()
        )
        self.spliter = RecursiveCharacterTextSplitter(
            separators=settings.chroma.separatoes,
            chunk_size=settings.chroma.chunk_size,
            chunk_overlap=settings.chroma.chunk_overlap,
            length_function=len,
        )
        self.law_splitter = SplitterFactory.legal_document_loader

    def get_retriever(self, category: str | None = None):
        search_kwargs = {"k": settings.chroma.k}
        if category:
            search_kwargs["filter"] = {"category": category}
        return self.vector_store.as_retriever(search_kwargs=search_kwargs)

    def _resolve_category(self, file_path: Path) -> str:
        """根据文件路径推断分类。优先使用相对目录，其次使用文件名。"""
        data_root = Path(settings.chroma.data_path)
        try:
            relative_path = file_path.relative_to(data_root)
        except ValueError:
            return file_path.stem

        if relative_path.parent != Path("."):
            # 目录结构可表达业务类别时，优先按目录归类
            return str(relative_path.parent).replace("\\", "/")
        return file_path.stem

    def _enrich_documents_with_category(self, documents, category: str, source_path: Path):
        for doc in documents:
            base_metadata = dict(doc.metadata or {})
            base_metadata["category"] = category
            base_metadata.setdefault("source", str(source_path))
            doc.metadata = base_metadata
        return documents

    def _remove_md5_record(self, file_md5: str) -> bool:
        """从 MD5 存储文件中移除指定文件哈希。"""
        if not file_md5:
            return False

        md5_store_path = Path(settings.chroma.md5_hex_store)
        if not md5_store_path.exists() or not md5_store_path.is_file():
            return False

        lines = md5_store_path.read_text(encoding="utf-8").splitlines()
        new_lines = [line for line in lines if line.strip().lower() != file_md5.lower()]
        removed = len(new_lines) != len(lines)
        md5_store_path.write_text(
            "\n".join(new_lines) + ("\n" if new_lines else ""),
            encoding="utf-8",
        )
        return removed

    def remove_document_by_file(self, file_path: Path | str) -> dict:
        """删除指定文件在向量库中的分块，并移除对应 MD5 记录。"""
        target_path = Path(file_path)
        if not target_path.is_absolute():
            target_path = Path(settings.chroma.data_path) / target_path
        target_path = target_path.resolve()

        source = str(target_path)
        result = self.vector_store.get(where={"source": source}, include=["metadatas"])
        ids = result.get("ids", []) or []

        deleted_chunks = 0
        if ids:
            self.vector_store.delete(ids=ids)
            deleted_chunks = len(ids)

        file_md5 = get_file_md5_hex(target_path) if target_path.exists() else None
        md5_removed = self._remove_md5_record(file_md5)

        logger.info(
            f"删除文件索引完成: {target_path} | deleted_chunks={deleted_chunks} | md5_removed={md5_removed}"
        )
        return {
            "source": source,
            "deleted_chunks": deleted_chunks,
            "md5_removed": md5_removed,
        }

    def load_document(self):
        """
        从数据文件夹内读取数据文件，转为向量存入向量库
        要计算文件的MD5做去重
        :return: None
        """
        def load_docx_with_metadata(file_path: Path):
            # 使用 python-docx 读取，避免 unstructured 在部分环境触发额外 NLP 依赖
            doc = DocxDocument(str(file_path))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            content = "\n".join(paragraphs)
            if not content:
                return []
            return [Document(page_content=content, metadata={"source": str(file_path)})]
        def check_md5(md5_for_check: str) -> bool:
            if not md5_for_check:
                logger.warning("未配置 chroma.md5_hex_store，跳过 MD5 校验")
                return True # 或者根据业务逻辑返回 False
            return check_md5_hex(abs_path=Path(settings.chroma.md5_hex_store), md5_for_check=md5_for_check)
        def save_md5(md5_for_check: str):
            if not md5_for_check:
                logger.warning("未配置 chroma.md5_hex_store，跳过 MD5 校验")
                return True # 或者根据业务逻辑返回 False
            return save_md5_hex(abs_path=Path(settings.chroma.md5_hex_store), md5_for_check=md5_for_check)
        def get_file_documents(file_path: Path):
            # 这里你需要根据文件类型来加载文档，这只是一个示例
            suffix = file_path.suffix.lower()
            if suffix == ".pdf":
                return pdf_loader(file_path, passwd=None)
            elif suffix == ".txt":
                return text_loader(file_path)
            elif suffix == ".docx":
                return load_docx_with_metadata(file_path)
            else:
                logger.info(f"不支持的文件类型: {file_path.suffix}")
                raise FileTypeError(file_path.suffix)
        allowed_files_path: list = listdir_with_allowed_types(Path(settings.chroma.data_path), settings.chroma.allow_knowledge_file_type)
        for path in allowed_files_path:
            md5_hex = get_file_md5_hex(path)
            if check_md5(md5_hex):
                logger.info(f"文件已存在，跳过处理: {path}")
                continue
            try:
                documents = get_file_documents(path)
                if not documents:
                    logger.warning(f"未能从文件中提取到任何文档: {path}")
                    continue

                category = self._resolve_category(path)
                if category == "law":
                    source_text = "\n".join(doc.page_content for doc in documents if doc.page_content)
                    split_documents = self.law_splitter(
                        text=source_text,
                        law_name=path.stem,
                        source=str(path),
                    )
                else:
                    split_documents = self.spliter.split_documents(documents)

                if not split_documents:
                    logger.warning(f"未能从文件中提取到任何分块文档: {path}")
                    continue
                split_documents = self._enrich_documents_with_category(split_documents, category, path)
                self.vector_store.add_documents(split_documents)
                save_md5(md5_hex)
                logger.info(f"成功处理文件并添加到向量库: {path} | category={category}")
            except Exception as e:
                logger.warning(str(e))
if __name__ == '__main__': 
    vs = VectorStoreService()
    # vs.remove_document_by_file(Path("law/中华人民共和国劳动合同法_20121228.docx"))
    vs.load_document()